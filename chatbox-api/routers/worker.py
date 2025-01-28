import asyncio
import os
from langchain_openai import AzureChatOpenAI
from langchain_core.output_parsers import StrOutputParser
from langchain_core.prompts import ChatPromptTemplate
from langchain_openai import AzureOpenAIEmbeddings
from langchain_community.document_loaders import PyPDFLoader
from langchain_core.documents  import Document
from docx import Document as DocxDocument
from langchain_text_splitters import CharacterTextSplitter
from dotenv import load_dotenv
from pathlib import Path
from fastapi import APIRouter, HTTPException, WebSocket, WebSocketDisconnect,File, UploadFile
from langdetect import detect
from models.question import Question
from langchain_chroma import Chroma
from uuid import uuid4
import logging

logging.basicConfig(level=logging.DEBUG)

# Initialize FastAPI app
router = APIRouter()

# Load environment variables
load_dotenv()

# Azure OpenAI configuration
AZURE_API_KEY = os.getenv("AZURE_OPENAI_API_KEY")
AZURE_ENDPOINT = os.getenv("AZURE_ENDPOINT")
AZURE_API_VERSION = os.getenv("AZURE_API_VERSION")
AZURE_EMBEDDING_ENDPOINT = os.getenv("AZURE_EMBEDDING_ENDPOINT")

# Initialize AzureChatOpenAI for gpt-4o-mini
chat_model = AzureChatOpenAI(
    azure_deployment="gpt-4o-mini",
    api_key=AZURE_API_KEY,
    azure_endpoint=AZURE_ENDPOINT,
    api_version=AZURE_API_VERSION
)

UPLOAD_DIR = Path("./data_warehouse")
UPLOAD_DIR.mkdir(parents=True, exist_ok=True)

vector_store = None

async def ensure_vector_store_initialized():
    global vector_store
    folder = "./data_warehouse"
    embeddings = AzureOpenAIEmbeddings(
        model="text-embedding-ada-002",
        azure_endpoint=AZURE_EMBEDDING_ENDPOINT,
        api_key=AZURE_API_KEY,
        openai_api_version=AZURE_API_VERSION
    )
    vector_store = Chroma(
        collection_name="tickets_info",
        embedding_function=embeddings,
        persist_directory=folder,
    )

def load_docx(file_path):
    """
    Custom function to load content from a .docx file and return it as LangChain Document objects.
    """
    doc = DocxDocument(file_path)
    content = []
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():  # Ignore empty paragraphs
            content.append(paragraph.text.strip())

    # Return as a list of LangChain Document objects
    return [Document(page_content="\n".join(content), metadata={"source": str(file_path)})]

async def update_vector_store_with_file(file_path):
    """
    Load PDFs and DOCX files from a folder and create embeddings.
    """
    await ensure_vector_store_initialized()
    file_extension = file_path.suffix.lower()
    if file_extension == ".pdf":
        loader = PyPDFLoader(str(file_path))
        documents = loader.load()
    elif file_extension == ".docx":
        documents = load_docx(str(file_path))
    else:
        raise ValueError("Unsupported file type.")

    text_splitter = CharacterTextSplitter(chunk_size=1000, chunk_overlap=100)

    split_docs = text_splitter.split_documents(documents)
    uuids = [str(uuid4()) for _ in range(len(split_docs))]
    if split_docs:
        vector_store.add_documents(documents=split_docs, ids=uuids)
    else:
        raise Exception("No docs")

    return vector_store

async def find_closest_doc(prompt):
    """
    Search for the closest doc to a user's prompt.
    """
    await ensure_vector_store_initialized()
    try:
        if(len(vector_store.get()['documents'])>0):        
            result = vector_store.similarity_search(prompt,k=1)
            return result[0] if result else None
        else: 
            return {"error": f"Number of docs = {len(vector_store.get()['documents'])}"}
    except Exception as e:
        return {"error": f"{str(e)}"}

def summarize_pdf_with_gpt4o(question, text, detected_language):
    """
    Summarize a PDF's content using Azure OpenAI (gpt-4o-mini).
    """
    prompt = ChatPromptTemplate.from_template("""
            You are helping our company's employees with self-service. Given this information:

            {content}

            Answer: {question} in {language}.
        """)
    
    content_chain = prompt | chat_model

    response = content_chain.invoke({"content": text, "question": question, "language": detected_language})
    
    return response.content

async def process_user_prompt(prompt, pdf_folder, detected_language):
    """
    Full pipeline: find the closest PDF, extract content, and summarize.
    """
    try:
        closest_doc = await find_closest_doc(prompt)
        if closest_doc:
            pdf_content = closest_doc.page_content
            summary = summarize_pdf_with_gpt4o(prompt, pdf_content, detected_language)

            return {
                "closest_pdf": closest_doc.metadata["source"],
                "summary": summary
            }

        else:
            return {
                "error": "No relevant PDF found."
            }
    except Exception as e:
        return {"error": f"{str(e)}"}

@router.post("/ask")
async def ask_question(payload: Question):
    """
    API to process user question and return summary of the most relevant PDF.
    """
    pdf_folder = "./data_warehouse"
    question = payload.question

    if not question:
        raise HTTPException(status_code=400, detail="Question cannot be empty.")

    detected_language = detect(question)

    if detected_language != "vi":
        detected_language = "en"

    result = await process_user_prompt(question, pdf_folder, detected_language)

    if not result:
        raise HTTPException(status_code=404, detail="No relevant document found.")

    if "error" in result:
        raise HTTPException(status_code=404, detail=result["error"])

    if detected_language != "vi":
        answer = result["summary"]
    else:
        answer = result["summary"]
    return {
        "closest_pdf": result["closest_pdf"],
        "answer": answer
    }

@router.post("/upload")
async def upload_file(file: UploadFile = File(...)):
    """
    Endpoint to upload a file and update the vector_store.
    """
    file_extension = file.filename.split('.')[-1]
    if file_extension not in ["pdf", "docx"]:
        return {"error": "File type not allowed. Only .pdf and .docx are supported."}

    file_path = UPLOAD_DIR / file.filename
    
    if file_path.exists():
        return {"error": f"File {file.filename} already exists."}

    with open(file_path, "wb") as f:
        f.write(await file.read())

    try:
        task = asyncio.create_task(update_vector_store_with_file(file_path))
        await task  # Await the task to ensure it's completed
    except Exception as e:
        return {"error": f"Failed to update vector store: {str(e)}"}

    return {"message": f"File {file.filename} uploaded and vector store updated successfully!"}


async def stream_lines(data):
    lines = data.split("\n")
    for line in lines:
        yield line
        await asyncio.sleep(1)

@router.websocket("/ws")
async def websocket_endpoint(websocket: WebSocket):
    await websocket.accept()
    try:
        # Continuously listen for messages from the client
        while True:
            # Receive the payload from the client
            payload = await websocket.receive_text()
            print(f"Received payload: {payload}")
            
            pdf_folder = "./data_warehouse"
            question = payload.question
            if not question.strip():
                raise HTTPException(status_code=400, detail="Question cannot be empty.")

            result = process_user_prompt(question, pdf_folder)

            # Process the payload and stream the response
            if result: 
                async for line in stream_lines(result["summary"]):
                    await websocket.send_text(line)
    except WebSocketDisconnect:
        print("Client disconnected")

@router.on_event("startup")
async def initialize_vector_store():
    global vector_store
    await ensure_vector_store_initialized()
    print("Vector store initialized!")
